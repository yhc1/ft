"""
Generate two .docx files from dimensions.json and prerequisites.json:
  - AIL_Readiness_Assessment_EN.docx (English)
  - AIL_Readiness_Assessment_ZH.docx (Chinese)
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DATA_DIR = Path(__file__).parent / "web" / "data"
OUTPUT_DIR = Path(__file__).parent


def load_data():
    with open(DATA_DIR / "dimensions.json", encoding="utf-8") as f:
        dims_data = json.load(f)
    with open(DATA_DIR / "prerequisites.json", encoding="utf-8") as f:
        prereqs_data = json.load(f)
    return dims_data, prereqs_data


def set_heading_style(paragraph, level, color_hex="1F3864"):
    """Apply heading style with colour."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    sizes = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 11))


def add_heading(doc, text, level):
    heading_colors = {1: "1F3864", 2: "2E5FA3", 3: "2E5FA3"}
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.bold = True
        color = heading_colors.get(level, "000000")
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))
    return p


def add_label_value(doc, label, value, indent=False):
    """Add a paragraph with bold label followed by normal value."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_body(doc, text, indent=False, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_bullet(doc, text, indent_level=1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * indent_level)
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)


def build_document(lang: str, dims_data: dict, prereqs_data: dict) -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title_text = {
        "en": "AIL Readiness Assessment",
        "zh": "AIL 準備度評估",
    }[lang]
    title_p = doc.add_heading(title_text, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    # ── Foundational Requirements ────────────────────────────────────────────
    section_labels = {
        "en": {
            "foundational": "Foundational Requirements",
            "prerequisites": "Prerequisites",
            "use_cases": "Use Cases",
            "tech_assessments": "Technical Assessments",
            "modelling_process": "Modelling Process",
            "model_output": "Model Output",
            "reference_case": "Reference Case",
        },
        "zh": {
            "foundational": "基礎前提條件",
            "prerequisites": "前置資料需求",
            "use_cases": "應用場景",
            "tech_assessments": "技術評估",
            "modelling_process": "建模流程",
            "model_output": "模型輸出",
            "reference_case": "參考案例",
        },
    }[lang]

    add_heading(doc, section_labels["foundational"], level=1)

    for req in dims_data.get("foundationalRequirements", []):
        name = req["name"][lang]
        desc = req["description"][lang]
        p = doc.add_paragraph()
        name_run = p.add_run(f"{name}. ")
        name_run.bold = True
        name_run.font.size = Pt(10)
        desc_run = p.add_run(desc)
        desc_run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    add_separator(doc)

    # ── Dimensions ───────────────────────────────────────────────────────────
    for i, dim in enumerate(dims_data["dimensions"], start=1):
        dim_name = dim["name"][lang]
        add_heading(doc, f"{i}. {dim_name}", level=1)

        # Description
        add_body(doc, dim["description"][lang])

        # Prerequisites
        add_heading(doc, section_labels["prerequisites"], level=2)
        for prereq in dim.get("prerequisites", []):
            display = prereq["displayName"][lang]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            label_run = p.add_run(f"{display}")
            label_run.bold = True
            label_run.font.size = Pt(10)
            for item in prereq.get("items", []):
                add_bullet(doc, item[lang])

        # Use Cases
        add_heading(doc, section_labels["use_cases"], level=2)
        for uc in dim.get("useCases", []):
            uc_name = uc["name"][lang]
            add_heading(doc, uc_name, level=3)
            add_label_value(doc, section_labels["modelling_process"], uc["modellingProcess"][lang], indent=True)
            add_label_value(doc, section_labels["model_output"], uc["modelOutput"][lang], indent=True)
            ref = uc.get("referenceCase", "—") or "—"
            add_label_value(doc, section_labels["reference_case"], ref, indent=True)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Technical Assessments
        add_heading(doc, section_labels["tech_assessments"], level=2)
        for assessment in dim.get("technicalAssessments", []):
            title = assessment["title"][lang]
            content = assessment["content"][lang]
            p = doc.add_paragraph()
            title_run = p.add_run(f"{title}: ")
            title_run.bold = True
            title_run.font.size = Pt(10)
            content_run = p.add_run(content)
            content_run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        # Separator between dimensions
        if i < len(dims_data["dimensions"]):
            doc.add_page_break()

    return doc


def main():
    dims_data, prereqs_data = load_data()

    for lang, suffix in [("en", "EN"), ("zh", "ZH")]:
        doc = build_document(lang, dims_data, prereqs_data)
        out_path = OUTPUT_DIR / f"AIL_Readiness_Assessment_{suffix}.docx"
        doc.save(out_path)
        print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
